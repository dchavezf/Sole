"""
DOCX Exporter: translates the markdown-it token stream directly to a
python-docx Document, reading all style properties from the theme JSON.

Images are always embedded (base64 bytes in the .docx zip).
Callouts and directives use paragraph borders (w:pBdr) and shading (w:shd).
"""

from __future__ import annotations

import io
import urllib.request
import warnings
from pathlib import Path
from typing import Any

from markdown_it.token import Token

from core.theme import Theme
from core.ast_walker import (
    iter_blocks, Block,
    is_heading, heading_level,
    is_callout, callout_type,
    is_directive, directive_name,
    is_table,
    get_inline_text,
    first_inline,
)
from core.field_extractor import extract

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu, Twips, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

def _trim_checkbox_children(children: list) -> list:
    """Remove the leading '[ ] ' or '[x] ' text from inline children of a task list item."""
    import re
    _CB_RE = re.compile(r"^\[[ xX]\]\s*")
    result = list(children)
    for i, child in enumerate(result):
        if child.type == "text" and _CB_RE.match(child.content):
            trimmed = _CB_RE.sub("", child.content)
            if trimmed:
                child.content = trimmed
            else:
                result.pop(i)
            break
    return result


_ALIGN_MAP = {
    "left":     WD_ALIGN_PARAGRAPH.LEFT      if _HAS_DOCX else None,
    "center":   WD_ALIGN_PARAGRAPH.CENTER    if _HAS_DOCX else None,
    "right":    WD_ALIGN_PARAGRAPH.RIGHT     if _HAS_DOCX else None,
    "justify":  WD_ALIGN_PARAGRAPH.JUSTIFY   if _HAS_DOCX else None,
}

_CHECKBOX_UNCHECKED = "☐"
_CHECKBOX_CHECKED   = "☑"

class DocxExporter:
    def __init__(self, theme: Theme, base_dir: Path | None = None) -> None:
        self.theme = theme
        self.base_dir = base_dir
        self._h1_count = 0  # track to skip page break before first H1

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def export(self, tokens: list[Token], output_path: Path, title: str = "Documento") -> None:
        if not _HAS_DOCX:
            warnings.warn(
                "python-docx is not installed — DOCX export skipped.\n"
                "Install with: pip install python-docx",
                stacklevel=2,
            )
            return

        self._h1_count = 0
        doc = Document()
        self._configure_styles(doc)
        self._setup_page_layout(doc, title)
        self._render(doc, tokens)
        doc.save(str(output_path))

    # ------------------------------------------------------------------
    # Style configuration
    # ------------------------------------------------------------------

    def _configure_styles(self, doc: "Document") -> None:
        """Pre-configure named paragraph/character styles from the theme."""
        normal = doc.styles["Normal"]
        para_cfg = self.theme.get_element("paragraph")
        _apply_docx_run(normal.font, para_cfg.get("docx", {}).get("run", {}))

    # ------------------------------------------------------------------
    # Token rendering
    # ------------------------------------------------------------------

    def _render(self, doc: "Document", tokens: list[Token]) -> None:
        for item in iter_blocks(tokens):
            self._render_item(doc, item)

    def _render_item(self, doc: "Document", item: Block | Token) -> None:
        if isinstance(item, Token):
            self._render_token(doc, item)
        else:
            self._render_block(doc, item)

    def _render_token(self, doc: "Document", tok: Token) -> None:
        if tok.type == "fence":
            self._render_code_block(doc, tok)
        elif tok.type == "hr":
            self._render_hr(doc)
        elif tok.type == "pagebreak":
            self._render_pagebreak(doc)
        elif tok.type == "inline":
            p = doc.add_paragraph()
            self._render_inline(p, tok.content)

    def _render_block(self, doc: "Document", block: Block) -> None:
        open_type = block.open.type

        if open_type == "heading_open":
            self._render_heading(doc, block)
        elif open_type in ("paragraph_open",):
            self._render_paragraph(doc, block)
        elif open_type == "callout_open":
            self._render_callout(doc, block)
        elif open_type in ("container_open", "container_directive_open"):
            self._render_directive(doc, block)
        elif open_type == "table_open":
            self._render_table(doc, block)
        elif open_type == "bullet_list_open":
            self._render_list(doc, block, ordered=False)
        elif open_type == "ordered_list_open":
            self._render_list(doc, block, ordered=True)
        elif open_type == "blockquote_open":
            self._render_blockquote(doc, block)
        else:
            # Recurse into unknown blocks
            for child in block.children:
                self._render_item(doc, child)

    # ------------------------------------------------------------------
    # Element renderers
    # ------------------------------------------------------------------

    def _render_heading(self, doc: "Document", block: Block) -> None:
        level = heading_level(block.open)
        inline = first_inline(block)
        cfg = self.theme.get_heading(level)
        docx_cfg = cfg.get("docx", {})

        try:
            p = doc.add_heading("", level=level)
        except Exception:
            p = doc.add_paragraph()

        para_cfg = dict(docx_cfg.get("paragraph", {}))
        if level == 1:
            if self._h1_count > 0 and self.theme.get_layout().get("autoPageBreakOnH1", False):
                if not (block.open.meta and block.open.meta.get("noPageBreak")):
                    para_cfg["pageBreakBefore"] = True
            self._h1_count += 1

        _apply_paragraph_cfg(p, para_cfg)

        if inline:
            run_cfg = docx_cfg.get("run", {})
            self._render_inline_token(p, inline, extra_run_cfg=run_cfg)


    def _render_paragraph(self, doc: "Document", block: Block) -> None:
        inline = first_inline(block)
        if not inline:
            return
        cfg = self.theme.get_element("paragraph")
        docx_cfg = cfg.get("docx", {})
        p = doc.add_paragraph()
        _apply_paragraph_cfg(p, docx_cfg.get("paragraph", {}))
        self._render_inline_token(p, inline)

    def _render_code_block(self, doc: "Document", tok: Token) -> None:
        cfg = self.theme.get_element("codeBlock")
        docx_cfg = cfg.get("code", {}).get("docx", {}) if isinstance(cfg, dict) else {}
        p = doc.add_paragraph(tok.content.rstrip())
        style_name = docx_cfg.get("style", "")
        if style_name and style_name in [s.name for s in doc.styles]:
            p.style = style_name
        shading_cfg = docx_cfg.get("paragraph", {}).get("shading", {})
        if shading_cfg:
            _apply_shading(p, shading_cfg)
        run_cfg = docx_cfg.get("run", self.theme.get_element("inlineCode").get("docx", {}))
        for run in p.runs:
            _apply_docx_run(run.font, run_cfg)

    def _render_pagebreak(self, doc: "Document") -> None:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _render_hr(self, doc: "Document") -> None:
        p = doc.add_paragraph()
        cfg = self.theme.get_element("horizontalRule")
        _apply_paragraph_border(p, cfg.get("docx", {}).get("paragraph", {}).get("border", {}))

    def _render_callout(self, doc: "Document", block: Block) -> None:
        ct = callout_type(block)
        cfg = self.theme.get_callout(ct)
        docx_cfg = cfg.get("docx", {})
        label = cfg.get("label") or ""

        shading = docx_cfg.get("paragraph", {}).get("shading", {})

        if label:
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(f"  {label.upper()}")
            title_run.bold = True
            _apply_callout_border(title_p, docx_cfg, is_title=True)
            if shading:
                _apply_shading(title_p, shading)

        for child in block.children:
            p = doc.add_paragraph()
            if isinstance(child, Token) and child.type == "inline":
                self._render_inline_token(p, child)
            elif isinstance(child, Block):
                inner = first_inline(child)
                if inner:
                    self._render_inline_token(p, inner)
            _apply_callout_border(p, docx_cfg, is_title=False)
            if shading:
                _apply_shading(p, shading)

    def _render_directive(self, doc: "Document", block: Block) -> None:
        name = directive_name(block)
        cfg = self.theme.get_directive(name)
        docx_cfg = cfg.get("docx", {})
        fields_schema = cfg.get("fields", {})

        if fields_schema and docx_cfg.get("fieldSequence"):
            self._render_directive_with_fields(doc, block, cfg)
            return

        # Generic fallback
        label = cfg.get("label", "")
        if label:
            lp = doc.add_paragraph()
            lp.add_run(label).bold = True
            _apply_callout_border(lp, docx_cfg, is_title=True)

        for child in block.children:
            p = doc.add_paragraph()
            if isinstance(child, Token) and child.type == "inline":
                self._render_inline_token(p, child)
            elif isinstance(child, Block):
                inner = first_inline(child)
                if inner:
                    self._render_inline_token(p, inner)
            shading = docx_cfg.get("paragraph", {}).get("shading", {})
            if shading:
                _apply_shading(p, shading)
            _apply_callout_border(p, docx_cfg, is_title=False)

    def _render_directive_with_fields(
        self, doc: "Document", block: Block, cfg: dict
    ) -> None:
        result = extract(block, cfg.get("fields", {}))
        fields = result.fields
        unmatched = result.unmatched

        docx_cfg = cfg.get("docx", {})
        para_cfg = docx_cfg.get("paragraph", {})
        field_sequence = docx_cfg.get("fieldSequence", [])

        for seq_item in field_sequence:
            field_key = seq_item["field"]
            render_as = seq_item.get("renderAs", "paragraph")

            if field_key == "_body":
                for child in unmatched:
                    p = doc.add_paragraph()
                    if isinstance(child, Token) and child.type == "inline":
                        self._render_inline_token(p, child)
                    elif isinstance(child, Block):
                        inner = first_inline(child)
                        if inner:
                            self._render_inline_token(p, inner)
                    if para_cfg:
                        _apply_paragraph_cfg(p, para_cfg)
                continue

            value = fields.get(field_key)
            if not value:
                continue

            if render_as == "heading":
                level = seq_item.get("level", 2)
                try:
                    p = doc.add_heading("", level=level)
                except Exception:
                    p = doc.add_paragraph()
                p.add_run(str(value)).bold = True

            elif render_as == "bold-prefix":
                prefix = seq_item.get("prefix", field_key)
                p = doc.add_paragraph()
                p.add_run(f"{prefix}: ").bold = True
                p.add_run(str(value))
                if para_cfg:
                    _apply_paragraph_cfg(p, para_cfg)

            elif render_as == "labeled-list":
                for entry in (value if isinstance(value, list) else []):
                    p = doc.add_paragraph()
                    p.add_run(f"{entry['label']}: ").bold = True
                    p.add_run(entry["value"])
                    if para_cfg:
                        _apply_paragraph_cfg(p, para_cfg)

            elif render_as == "paragraph":
                p = doc.add_paragraph(str(value))
                if para_cfg:
                    _apply_paragraph_cfg(p, para_cfg)

    def _render_blockquote(self, doc: "Document", block: Block) -> None:
        for child in block.children:
            p = doc.add_paragraph()
            if isinstance(child, Token) and child.type == "inline":
                self._render_inline_token(p, child)
            elif isinstance(child, Block):
                inner = first_inline(child)
                if inner:
                    self._render_inline_token(p, inner)
            cfg = self.theme.get_element("blockquote")
            _apply_paragraph_cfg(p, cfg.get("docx", {}).get("paragraph", {}))

    def _render_list(self, doc: "Document", block: Block, ordered: bool) -> None:
        list_style = "List Number" if ordered else "List Bullet"
        for item in block.children:
            if not isinstance(item, Block):
                continue
            inline = first_inline(item)
            content = get_inline_text(inline) if inline else ""
            is_checked = content.startswith("[x]")
            is_task = content.startswith("[ ]") or is_checked
            checkbox_cfg = self.theme.get_element("taskList").get("checkbox", {}).get("docx", {})

            p = doc.add_paragraph(style=list_style)
            if is_task:
                symbol = checkbox_cfg.get("checkedSymbol") if is_checked else checkbox_cfg.get("uncheckedSymbol")
                symbol_run = p.add_run(f"{symbol or ''} ")
                if checkbox_cfg.get("font"):
                    symbol_run.font.name = checkbox_cfg["font"]
                # strip the checkbox prefix and render remaining inline tokens
                if inline and inline.children:
                    # skip first text child that is "[ ] " or "[x] "
                    trimmed = _trim_checkbox_children(inline.children)
                    self._render_inline_token(p, inline, override_children=trimmed)
                else:
                    p.add_run(content[3:].strip())
            elif inline:
                self._render_inline_token(p, inline)

    def _render_table(self, doc: "Document", block: Block) -> None:
        # Collect (inline_token | None) per cell
        rows: list[list[Token | None]] = []
        header_count = 0

        for section in block.children:
            if not isinstance(section, Block):
                continue
            is_header = section.open.type == "thead_open"
            for row_block in section.children:
                if not isinstance(row_block, Block):
                    continue
                cells: list[Token | None] = []
                for cell in row_block.children:
                    if isinstance(cell, Block):
                        cells.append(first_inline(cell))
                    else:
                        cells.append(None)
                rows.append(cells)
                if is_header:
                    header_count += 1

        if not rows:
            return

        col_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"

        col_widths = self.theme.get_table_columns()
        for ri, row in enumerate(rows):
            tr = table.rows[ri]
            for ci, cell_tok in enumerate(row):
                cell = tr.cells[ci]
                # Clear default empty paragraph, then render inline
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                if cell_tok:
                    self._render_inline_token(p, cell_tok,
                                              extra_run_cfg={"bold": True} if ri < header_count else {})
                # Column width from theme
                cw_cfg = col_widths.get(str(ci), {})
                docx_w = cw_cfg.get("docxWidth")
                if docx_w:
                    cell.width = Twips(docx_w)

    # ------------------------------------------------------------------
    # Page layout: margins, headers, footers
    # ------------------------------------------------------------------

    def _setup_page_layout(self, doc: "Document", title: str) -> None:
        layout = self.theme.get_layout()
        meta = self.theme.meta
        docx_styles = self.theme.get_styles().get("docx", {})

        def _resolve_text(tmpl: str) -> str:
            return (tmpl
                    .replace("{{meta.author}}", meta.get("author", ""))
                    .replace("{{title}}", title))

        section = doc.sections[0]

        # Margins
        m = layout.get("margins", {}).get("docx", {})
        if m.get("top") is not None:
            section.top_margin = Cm(m["top"])
        if m.get("bottom") is not None:
            section.bottom_margin = Cm(m["bottom"])
        if m.get("left") is not None:
            section.left_margin = Cm(m["left"])
        if m.get("right") is not None:
            section.right_margin = Cm(m["right"])

        # Header
        h_cfg = layout.get("header", {})
        if h_cfg.get("enabled", True):
            header = section.header
            hp = header.paragraphs[0]
            hp.clear()

            left_text  = _resolve_text(h_cfg.get("left",   ""))
            center_text = _resolve_text(h_cfg.get("center", ""))
            right_text = _resolve_text(h_cfg.get("right",  ""))

            if left_text and right_text:
                # left + tab + right using a center tab stop at midpoint
                hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r1 = hp.add_run(left_text)
                _apply_docx_run(r1.font, docx_styles.get("headerRun", {}))
                hp.add_run("\t")
                r2 = hp.add_run(right_text)
                _apply_docx_run(r2.font, docx_styles.get("headerRightRun", docx_styles.get("headerRun", {})))
                # right tab stop at page width
                from docx.oxml import OxmlElement as Ox
                pPr2 = hp._p.get_or_add_pPr()
                tabs = Ox("w:tabs")
                tab = Ox("w:tab")
                tab.set(qn("w:val"), "right")
                tab_pos = docx_styles.get("rightTabPos")
                if tab_pos is not None:
                    tab.set(qn("w:pos"), str(tab_pos))
                    tabs.append(tab); pPr2.append(tabs)
            elif center_text:
                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = hp.add_run(center_text)
                _apply_docx_run(r.font, docx_styles.get("headerRun", {}))
            elif right_text:
                hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = hp.add_run(right_text)
                _apply_docx_run(r.font, docx_styles.get("headerRightRun", docx_styles.get("headerRun", {})))
            elif left_text:
                r = hp.add_run(left_text)
                _apply_docx_run(r.font, docx_styles.get("headerRun", {}))

            if h_cfg.get("separator", True) and (left_text or center_text or right_text):
                _apply_paragraph_border(hp, docx_styles.get("headerSeparatorBorder", {}))

        # Footer
        f_cfg = layout.get("footer", {})
        if f_cfg.get("enabled", True):
            footer = section.footer
            fp = footer.paragraphs[0]
            fp.clear()

            _ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                      "center": WD_ALIGN_PARAGRAPH.CENTER,
                      "right": WD_ALIGN_PARAGRAPH.RIGHT}

            for pos in ("left", "center", "right"):
                val = f_cfg.get(pos, "")
                if not val:
                    continue
                fp.alignment = _ALIGN[pos]
                if val == "pageNumber":
                    _add_page_number_field(
                        fp,
                        docx_styles.get("pageNumberRun", docx_styles.get("footerRun", {})),
                        docx_styles.get("pageNumberSeparator", ""),
                    )
                else:
                    r = fp.add_run(_resolve_text(val))
                    _apply_docx_run(r.font, docx_styles.get("footerRun", {}))
                break  # only first non-empty position

    # ------------------------------------------------------------------
    # Inline rendering — walks markdown-it children (properly parsed tokens)
    # ------------------------------------------------------------------

    def _render_inline_token(
        self,
        paragraph: Any,
        tok: "Token",
        extra_run_cfg: dict | None = None,
        override_children: list | None = None,
    ) -> None:
        """
        Walk an inline token's children and add formatted runs to the paragraph.

        Children types handled:
          text, softbreak, hardbreak, strong_open/close, em_open/close,
          s_open/close, code_inline, link_open/close, html_inline (skipped)
        """
        bold = False
        italic = False
        strike = False
        link_href: str | None = None
        children = override_children if override_children is not None else (tok.children or [])

        # Fallback: if no children but content exists, add as plain run
        if not children and tok.content:
            run = paragraph.add_run(tok.content)
            if extra_run_cfg:
                _apply_docx_run(run.font, extra_run_cfg)
            return

        for child in children:
            ct = child.type
            if ct == "text":
                run = paragraph.add_run(child.content)
                run.bold = bold or None
                run.italic = italic or None
                run.font.strike = strike or None
                if link_href:
                    link_cfg = self.theme.get_element("link").get("docx", {}).get("run", {})
                    _apply_docx_run(run.font, link_cfg)
                if extra_run_cfg:
                    _apply_docx_run(run.font, extra_run_cfg)
            elif ct == "code_inline":
                run = paragraph.add_run(child.content)
                code_cfg = self.theme.get_element("inlineCode").get("docx", {})
                _apply_docx_run(run.font, code_cfg)
                if extra_run_cfg:
                    _apply_docx_run(run.font, extra_run_cfg)
            elif ct in ("softbreak", "hardbreak"):
                paragraph.add_run(" ")
            elif ct == "strong_open":
                bold = True
            elif ct == "strong_close":
                bold = False
            elif ct == "em_open":
                italic = True
            elif ct == "em_close":
                italic = False
            elif ct == "s_open":
                strike = True
            elif ct == "s_close":
                strike = False
            elif ct == "link_open":
                link_href = child.attrGet("href") or ""
            elif ct == "link_close":
                link_href = None
            # html_inline, image, fence — skip


# ------------------------------------------------------------------
# DOCX property helpers
# ------------------------------------------------------------------

def _add_page_number_field(paragraph: Any, run_cfg: dict | None = None, separator_text: str = "") -> None:
    """Insert PAGE / NUMPAGES field into a paragraph."""
    run_cfg = run_cfg or {}

    def _field_run(text: str, field: str) -> None:
        run = paragraph.add_run(text)
        _apply_docx_run(run.font, run_cfg)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {field} "
        run._r.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

    _field_run("", "PAGE")
    sep = paragraph.add_run(separator_text)
    _apply_docx_run(sep.font, run_cfg)
    _field_run("", "NUMPAGES")


def _hex_to_rgb(hex_color: str) -> RGBColor | None:
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 6:
        try:
            return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))
        except ValueError:
            return None
    return None


def _apply_docx_run(font: Any, run_cfg: dict) -> None:
    if not run_cfg:
        return
    if run_cfg.get("bold"):
        font.bold = True
    if run_cfg.get("italic"):
        font.italic = True
    if run_cfg.get("strike"):
        font.strike = True
    if run_cfg.get("underline"):
        font.underline = True
    if run_cfg.get("size"):
        font.size = Pt(run_cfg["size"] / 2)  # half-points → points
    if run_cfg.get("font"):
        font.name = run_cfg["font"]
    color = run_cfg.get("color")
    if color and color.startswith("#"):
        rgb = _hex_to_rgb(color)
        if rgb:
            font.color.rgb = rgb


def _apply_paragraph_cfg(paragraph: Any, para_cfg: dict) -> None:
    if not para_cfg:
        return
    fmt = paragraph.paragraph_format
    spacing = para_cfg.get("spacing", {})
    if spacing.get("before") is not None:
        fmt.space_before = Twips(spacing["before"])
    if spacing.get("after") is not None:
        fmt.space_after = Twips(spacing["after"])
    indent = para_cfg.get("indent", {})
    if indent.get("left") is not None:
        fmt.left_indent = Twips(indent["left"])
    align = para_cfg.get("alignment")
    if align and align in _ALIGN_MAP and _ALIGN_MAP[align]:
        paragraph.alignment = _ALIGN_MAP[align]
    if para_cfg.get("keepNext"):
        fmt.keep_with_next = True
    if para_cfg.get("keepLines"):
        fmt.keep_together = True
    if para_cfg.get("pageBreakBefore"):
        fmt.page_break_before = True
    _apply_paragraph_border(paragraph, para_cfg.get("border", {}))


def _apply_paragraph_border(paragraph: Any, border_cfg: dict) -> None:
    if not border_cfg:
        return

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    for side, cfg in border_cfg.items():
        if not isinstance(cfg, dict):
            continue
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"), str(cfg.get("style", "single")))
        node.set(qn("w:sz"), str(cfg.get("size", 4)))
        node.set(qn("w:space"), str(cfg.get("space", 1)))
        node.set(qn("w:color"), str(cfg.get("color", "")).lstrip("#").upper())
        existing = pBdr.find(qn(f"w:{side}"))
        if existing is not None:
            pBdr.remove(existing)
        pBdr.append(node)


def _apply_shading(paragraph: Any, shading_cfg: dict) -> None:
    fill = shading_cfg.get("fill", "")
    if not fill:
        return
    fill = fill.lstrip("#")
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), shading_cfg.get("type", "clear"))
    shd.set(qn("w:color"), str(shading_cfg.get("color", "auto")).lstrip("#"))
    shd.set(qn("w:fill"), fill.upper())
    existing = pPr.find(qn("w:shd"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(shd)


def _apply_callout_border(paragraph: Any, docx_cfg: dict, is_title: bool) -> None:
    """Apply left border to simulate callout styling using w:pBdr."""
    border_cfg = docx_cfg.get("paragraph", {}).get("border", {})
    left_border = border_cfg.get("left", {})
    if not left_border:
        return

    required = ("color", "size", "space", "style")
    if any(left_border.get(key) is None for key in required):
        return

    color = str(left_border["color"]).lstrip("#")
    size  = str(left_border["size"])
    space = str(left_border["space"])
    style = str(left_border["style"])

    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    else:
        pBdr = existing

    left = OxmlElement("w:left")
    left.set(qn("w:val"),   style)
    left.set(qn("w:sz"),    size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color.upper())
    existing_left = pBdr.find(qn("w:left"))
    if existing_left is not None:
        pBdr.remove(existing_left)
    pBdr.append(left)
